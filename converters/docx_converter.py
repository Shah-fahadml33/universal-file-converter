from pathlib import Path
import platform

from docx import Document



def docx_to_pdf(docx_path: str, output_path: str):
    """
    Convert DOCX to PDF.

    Note:
    This requires Microsoft Word on Windows/macOS.
    """

    if platform.system() not in ["Windows", "Darwin"]:
        raise RuntimeError(
            "DOCX to PDF using docx2pdf is only supported on Windows and macOS."
        )

    from docx2pdf import convert

    convert(docx_path, output_path)

    return output_path


def docx_to_txt(docx_path: str, output_path: str):
    """
    Convert DOCX to a plain text file.
    """

    document = Document(docx_path)

    paragraphs = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(paragraphs))

    return output_path


def extract_docx_info(docx_path: str):
    """
    Extract useful information from a DOCX file.
    """

    document = Document(docx_path)

    info = {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "words": 0,
        "characters": 0,
    }

    word_count = 0
    character_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text
        word_count += len(text.split())
        character_count += len(text)

    info["words"] = word_count
    info["characters"] = character_count

    return info


def get_docx_text(docx_path: str):
    """
    Return the complete text from a DOCX file.
    """

    document = Document(docx_path)

    lines = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)

    return "\n".join(lines)


def validate_docx(docx_path: str):
    """
    Check whether a DOCX file is valid.
    """

    try:
        Document(docx_path)
        return True
    except Exception:
        return False


def copy_docx(docx_path: str, destination_folder: str):
    """
    Copy a DOCX file to another folder.
    """

    destination = Path(destination_folder)
    destination.mkdir(parents=True, exist_ok=True)

    destination_file = destination / Path(docx_path).name

    with open(docx_path, "rb") as src:
        with open(destination_file, "wb") as dst:
            dst.write(src.read())

    return str(destination_file)